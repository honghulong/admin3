# -*- coding: utf-8 -*-
"""更新 mcp单点登录20260610.docx，补充 bind_employee 相关章节"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import copy

DOC_PATH = r'd:\hongmengProjects\admin3\mcp单点登录20260610.docx'

doc = Document(DOC_PATH)

# ============================================================
# 找到插入点：在 "3.4 XiaoYiAuthMcpConfig" 段落之后、"四、关键数据流" 之前
# ============================================================
insert_pos = None  # 在 insert_pos 之前插入
for i, p in enumerate(doc.paragraphs):
    if '四、关键数据流' in p.text:
        insert_pos = i
        break

# 找到 3.4 的末尾（3.4 的最后一个子段落）
last_of_34 = None
for i, p in enumerate(doc.paragraphs):
    if '3.4 XiaoYiAuthMcpConfig' in p.text:
        last_of_34 = i
    if last_of_34 is not None and i > last_of_34 and 'Heading 1' in p.style.name:
        break
    if '四、关键数据流' in p.text:
        break

print(f'Insert before paragraph {insert_pos}')

# ============================================================
# 构建要插入的内容
# ============================================================

def add_heading(doc, text, level):
    """添加标题"""
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, style='Normal', bold=False, bullet=False):
    """添加段落"""
    if bullet:
        p = doc.add_paragraph(style='List Bullet')
    else:
        p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    if bold:
        run.bold = True
    return p

def add_code_block(doc, text):
    """添加代码块（用等宽字体）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

# 构建新内容
new_elements = []

# 3.5 XiaoYiBindingMcpConfig
new_elements.append(('heading', '3.5 XiaoYiBindingMcpConfig（绑定服务端点）', 2))
new_elements.append(('para', '路径：infra/mcp/xiaoyi/XiaoYiBindingMcpConfig.java'))
new_elements.append(('para', '注册了一个独立的 MCP Server，端点路径为 /xiaoyi-mcp/binding/message。与 auth 服务分离，避免协议混淆。'))
new_elements.append(('para', '工具列表：'))
new_elements.append(('bullet', 'check_binding_status'))
new_elements.append(('bullet_indent', '描述：检查当前小艺会话的身份绑定状态'))
new_elements.append(('bullet_indent', '输入参数：无（agentLoginSessionId 由 Filter 从请求头自动提取）'))
new_elements.append(('bullet_indent', '返回：文本格式，首行为状态标识（bound/unbound/invalid）'))
new_elements.append(('bullet', 'bind_employee'))
new_elements.append(('bullet_indent', '描述：将当前小艺会话与指定的员工ID绑定'))
new_elements.append(('bullet_indent', '输入参数：xEmployeeId（String，必填），员工ID'))
new_elements.append(('bullet_indent', '返回：文本格式，绑定成功返回用户名和员工ID，失败返回错误提示'))

# 3.6 XiaoYiBindingService
new_elements.append(('heading', '3.6 XiaoYiBindingService（绑定逻辑）', 2))
new_elements.append(('para', '路径：infra/mcp/xiaoyi/XiaoYiBindingService.java'))
new_elements.append(('para', '核心方法：'))
new_elements.append(('bullet', 'bindEmployee(agentLoginSessionId, xEmployeeId)'))
new_elements.append(('bullet_indent', '输入：agentLoginSessionId（由 Filter 从 Header 提取）+ xEmployeeId（用户传入）'))
new_elements.append(('bullet_indent', '流程：'))
new_elements.append(('bullet_indent2', '1. 通过 userRepository.findByXEmployeeId() 核查员工ID是否存在'))
new_elements.append(('bullet_indent2', '2. 查找或创建 XiaoYiUserSession（session 不存在则新建，过期则报错）'))
new_elements.append(('bullet_indent2', '3. 将 User 关联到 XiaoYiUserSession，保存绑定关系'))
new_elements.append(('bullet_indent', '返回：BindEmployeeResult 记录类型'))
new_elements.append(('bullet_indent2', 'success（boolean）：是否绑定成功'))
new_elements.append(('bullet_indent2', 'username（String）：绑定的内部用户名'))
new_elements.append(('bullet_indent2', 'xEmployeeId（String）：绑定的员工ID'))
new_elements.append(('bullet_indent2', 'message（String）：成功或失败提示'))

# 3.7 绑定状态与员工关系
new_elements.append(('heading', '3.7 绑定状态与员工关系', 2))
new_elements.append(('para', '数据库表关系：'))
new_elements.append(('bullet', 'sys_user 表：已有 x_employee_id 字段（@Column(unique = true)），可在用户管理页面查看和修改'))
new_elements.append(('bullet', 'xiao_yi_user_session 表：通过 user_id 关联到 sys_user 表'))
new_elements.append(('bullet', '映射链路：agentLoginSessionId → XiaoYiUserSession.user → User.xEmployeeId'))
new_elements.append(('para', ''))
new_elements.append(('para', '绑定流程示意图：'))
new_elements.append(('code', '工作流调用 bind_employee(xEmployeeId="EMP001")'))
new_elements.append(('code', '  │'))
new_elements.append(('code', '  ├─ XiaoYiAuthFilter 从 Header 提取 agentLoginSessionId'))
new_elements.append(('code', '  ├─ XiaoYiBindingService.bindEmployee()'))
new_elements.append(('code', '  │  ├─ 1. 查 sys_user WHERE x_employee_id = "EMP001"'))
new_elements.append(('code', '  │  ├─ 2. 查/建 xiao_yi_user_session WHERE agent_login_session_id = ?'))
new_elements.append(('code', '  │  └─ 3. 设置 session.user = user，保存'))
new_elements.append(('code', '  └─ 返回绑定结果'))

# 更新 4.2 节（用户未绑定员工ID）
# 找到 4.2 节
section42_start = None
section42_end = None
for i, p in enumerate(doc.paragraphs):
    if '4.2 用户未绑定员工ID' in p.text:
        section42_start = i
    if section42_start is not None and i > section42_start and 'Heading' in p.style.name:
        section42_end = i
        break

print(f'4.2 section: paragraphs {section42_start} to {section42_end}')

# 更新 4.2 的内容（替换原有段落）
# 先删除 4.2 的原有内容（从 section42_start+1 到 section42_end）
# 由于 python-docx 删除段落比较麻烦，我们直接修改现有段落文本
if section42_start:
    # 修改第 1 行
    idx = section42_start + 1
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        p.clear()
        run = p.add_run('1. 工作流先调用 check_binding_status（无参数）')
        run.font.size = Pt(10.5)
    
    idx = section42_start + 2
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        p.clear()
        run = p.add_run('2. 返回 status=unbound，提示用户尚未绑定员工身份')
        run.font.size = Pt(10.5)
    
    idx = section42_start + 3
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        p.clear()
        run = p.add_run('3. 工作流进入绑定子流程，调用 bind_employee 工具，用户输入员工ID（如 "EMP001"）')
        run.font.size = Pt(10.5)
    
    idx = section42_start + 4
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        p.clear()
        run = p.add_run('4. XiaoYiAuthFilter 从 HTTP Header 提取 agentLoginSessionId，XiaoYiBindingService 核查员工ID是否存在，若存在则绑定 session 与 User 的关系')
        run.font.size = Pt(10.5)
    
    idx = section42_start + 5
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        p.clear()
        run = p.add_run('5. 绑定成功后，后续 check_binding_status 返回 status=bound，携带 username 和 xEmployeeId')
        run.font.size = Pt(10.5)

# ============================================================
# 在 "四、关键数据流" 之前插入新内容
# ============================================================
# python-docx 不支持在任意位置插入段落，需要用 XML 操作
# 获取 body 元素
body = doc.element.body

# 找到 "四、关键数据流" 的 XML 元素
section4_elem = None
for p_elem in body:
    if p_elem.tag == qn('w:p'):
        texts = p_elem.itertext()
        full_text = ''.join(texts)
        if '四、关键数据流' in full_text:
            section4_elem = p_elem
            break

if section4_elem is None:
    print('ERROR: Could not find "四、关键数据流" element')
    exit(1)

# 在 section4_elem 之前插入新元素
# 先创建临时文档来生成元素
temp_doc = Document()

for elem_type, text, *args in new_elements:
    if elem_type == 'heading':
        level = args[0] if args else 1
        p = temp_doc.add_heading(text, level=level)
    elif elem_type == 'para':
        p = temp_doc.add_paragraph(text)
        for run in p.runs:
            run.font.size = Pt(10.5)
    elif elem_type == 'bullet':
        p = temp_doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    elif elem_type == 'bullet_indent':
        p = temp_doc.add_paragraph(style='List Bullet 2')
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    elif elem_type == 'bullet_indent2':
        p = temp_doc.add_paragraph(style='List Bullet 2')
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        p.paragraph_format.left_indent = Inches(0.6)
    elif elem_type == 'code':
        p = temp_doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # 将临时文档中的段落移到主文档
    for child in list(temp_doc.element.body):
        body.insert(list(body).index(section4_elem), child)
    # 清空临时文档
    for child in list(temp_doc.element.body):
        temp_doc.element.body.remove(child)

# 更新文件清单表格（Table 1）
table = doc.tables[1]
# 添加新行
row = table.add_row()
cells = row.cells
cells[0].text = 'XiaoYiBindingMcpConfig.java'
cells[1].text = 'infra/mcp/xiaoyi/XiaoYiBindingMcpConfig.java'
cells[2].text = '新增'
row = table.add_row()
cells = row.cells
cells[0].text = 'XiaoYiBindingService.java'
cells[1].text = 'infra/mcp/xiaoyi/XiaoYiBindingService.java'
cells[2].text = '新增'

# 更新待办事项 - 标记 bind_employee 为已完成
for p in doc.paragraphs:
    if '开发 bind_employee MCP 工具' in p.text:
        p.clear()
        run = p.add_run('[已完成] 开发 bind_employee MCP 工具（XiaoYiBindingMcpConfig + XiaoYiBindingService）')
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        break

doc.save(DOC_PATH)
print('Document updated successfully!')

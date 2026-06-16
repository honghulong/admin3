"""更新 mcp单点登录20260610.docx：新增 x-request-id fallback 说明"""
import docx
from docx import Document

DOC_PATH = r'D:\hongmengProjects\admin3\mcp单点登录20260610.docx'

doc = Document(DOC_PATH)

def find_para(text_fragment, start=0):
    for i, p in enumerate(doc.paragraphs):
        if i < start: continue
        if text_fragment in p.text: return i
    return None

def insert_paragraph_after(para, text, style='Normal'):
    new_p = doc.add_paragraph(text, style=style)
    para._element.addnext(new_p._element)
    return new_p

def add_bullet_after(para, text):
    return insert_paragraph_after(para, text, 'List Bullet')

# ============================================================
# 1. 更新 3.8 节：增加 x-request-id fallback 说明
# ============================================================
idx = find_para('3.8 设备标识机制（v2 变更）')
if idx:
    # 找到该节最后一个 bullet
    end_idx = idx
    for j in range(idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].style.name.startswith('Heading 2'):
            break
        end_idx = j
    ref = doc.paragraphs[end_idx]
    add_bullet_after(ref, 'v2.1 新增 x-request-id header fallback：')
    new_p = add_bullet_after(ref, '背景：小艺在 bind_employee 等业务工具调用时，不传 deviceInfo/session 上下文，仅传业务参数')
    new_p = add_bullet_after(new_p, 'x-request-id header 格式为：sessionId&1-random-uuid（& 前是 sessionId）')
    new_p = add_bullet_after(new_p, 'XiaoYiAuthFilter 从 x-request-id 提取 sessionId 存入 request attribute')
    new_p = add_bullet_after(new_p, 'contextExtractor 将其透传到 McpTransportContext')
    new_p = add_bullet_after(new_p, '工具 handler 优先从 arguments 提取 sid/sessionId，取不到则回退到 x-request-id sessionId')
    new_p = add_bullet_after(new_p, '优先级：arguments.deviceInfo.sid > arguments.session.sessionId > x-request-id header')
elif idx is None:
    # 3.8 节不存在，在 3.7 后面创建
    idx37 = find_para('3.7 绑定状态与员工关系')
    if idx37:
        end37 = idx37
        for j in range(idx37 + 1, len(doc.paragraphs)):
            if doc.paragraphs[j].style.name.startswith('Heading'):
                end37 = j - 1
                break
        if end37 < idx37: end37 = idx37
        ref = doc.paragraphs[end37]
        new_h = insert_paragraph_after(ref, '3.8 设备标识机制（v2 变更）', 'Heading 2')
        add_bullet_after(new_h, '背景：小艺平台目前不传 agentLoginSessionId（因华为授权流程未开通），改为 MCP 请求 arguments 中传递设备标识，并通过 x-request-id header fallback')
        add_bullet_after(new_h, '设备标识来源（按优先级）：')
        add_bullet_after(new_h, '1) arguments.deviceInfo.sid（设备唯一标识，同一设备不变）')
        add_bullet_after(new_h, '2) arguments.session.sessionId（会话ID，每次会话不同）')
        add_bullet_after(new_h, '3) x-request-id header（sessionId&1-random，用于 bind_employee 等业务工具调用）')
        add_bullet_after(new_h, '数据库变更：xiao_yi_user_session 表新增 sid（唯一索引）和 session_id 列')
        p = add_bullet_after(new_h, '保留 agentLoginSessionId 结构，留待后续华为授权开通后使用')

# ============================================================
# 2. 更新 5. 设备标识传递机制
# ============================================================
idx5 = find_para('五、设备标识传递机制（v2 变更：从 agentLoginSessionId 改为 sid/sessionId）')
if idx5 is None:
    idx5 = find_para('五、设备标识传递机制')
if idx5:
    # 找到该节最后一个 bullet
    end5 = idx5
    for j in range(idx5 + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].style.name.startswith('Heading 1'):
            end5 = j - 1
            break
    if end5 < idx5: end5 = idx5
    
    ref = doc.paragraphs[end5]
    add_bullet_after(ref, 'v2.1 新增 x-request-id header fallback：')
    add_bullet_after(ref, 'XiaoYiAuthFilter 从 x-request-id header 提取 sessionId（& 前部分），存入 request attribute')
    p = add_bullet_after(ref, 'contextExtractor 将 request attribute 中的 sessionId 透传到 McpTransportContext')
    p = add_bullet_after(ref, '工具 handler 按优先级获取：arguments.deviceInfo.sid > arguments.session.sessionId > x-request-id')

# ============================================================
# 3. 更新文件清单表
# ============================================================
if len(doc.tables) >= 2:
    table = doc.tables[1]
    table.add_row()
    last_row = table.rows[-1]
    last_row.cells[0].text = 'XiaoYiAuthFilter.java'
    last_row.cells[1].text = 'infra/mcp/xiaoyi/XiaoYiAuthFilter.java'
    last_row.cells[2].text = '已修改：新增 x-request-id sessionId 提取'

# ============================================================
# 4. 更新待办事项
# ============================================================
idx_todo = find_para('八、待办事项')
if idx_todo:
    todo_items = []
    for j in range(idx_todo + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].style.name == 'List Bullet':
            if doc.paragraphs[j].text.strip():
                todo_items.append(j)
        elif doc.paragraphs[j].style.name.startswith('Heading') and j > idx_todo:
            break
    if todo_items:
        ref = doc.paragraphs[todo_items[-1]]
        add_bullet_after(ref, '[已完成] x-request-id header fallback：XiaoYiAuthFilter 从 x-request-id 提取 sessionId')
        add_bullet_after(ref, '测试 x-request-id fallback 绑定流程')

# Save
doc.save(DOC_PATH)
print('文档更新完成！')

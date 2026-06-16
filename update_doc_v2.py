import docx
from docx import Document
from docx.shared import Pt

DOC_PATH = r'D:\hongmengProjects\admin3\mcp单点登录20260610.docx'

doc = Document(DOC_PATH)

def find_para(text_fragment, start=0):
    for i, p in enumerate(doc.paragraphs):
        if i < start: continue
        if text_fragment in p.text: return i
    return None

def insert_paragraph_after(para, text, style='Normal'):
    new_p = doc.add_paragraph(text, style=style)
    para._element.addnext(new_p._element)
    return new_p

def add_bullet_after(para, text):
    return insert_paragraph_after(para, text, 'List Bullet')

# ============================================================
# 1. Add a new section "3.8 设备标识机制" after "3.7 绑定状态与员工关系"
# ============================================================
idx = find_para('3.7 绑定状态与员工关系')
if idx:
    # Find the last bullet under 3.7 (the end of this section)
    end_37 = idx
    for j in range(idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].style.name.startswith('Heading'):
            end_37 = j - 1
            break
    if end_37 < idx: end_37 = idx
    
    ref = doc.paragraphs[end_37]
    new_h = insert_paragraph_after(ref, '3.8 设备标识机制（v2 变更）', 'Heading 2')
    new_p = add_bullet_after(new_h, '背景：小艺平台目前不传 agentLoginSessionId（因华为授权流程未开通），改由 MCP 请求 body 中传递设备标识')
    new_p = add_bullet_after(new_p, '设备标识来源：')
    new_p = add_bullet_after(new_p, '优先使用 deviceInfo.sid（设备唯一标识，同一设备不变）')
    new_p = add_bullet_after(new_p, '其次使用 session.sessionId（会话ID，每次会话不同）')
    new_p = add_bullet_after(new_p, '服务端 XiaoYiAuthService.resolveDeviceKey(args) 从 arguments 中提取，sid 优先')
    new_p = add_bullet_after(new_p, '数据库变更：')
    new_p = add_bullet_after(new_p, 'xiao_yi_user_session 表新增 sid 列（唯一索引）和 session_id 列')
    new_p = add_bullet_after(new_p, 'agentLoginSessionId 改为非唯一（保留结构，但不用于验证）')
    new_p = add_bullet_after(new_p, '查找会话逻辑：先按 sid 查，找不到再按 sessionId 查')
    new_p = add_bullet_after(new_p, '保留 agentLoginSessionId 结构和 authorize/deauthorize 工具，留待后续华为授权开通后使用')

# ============================================================
# 2. Update tool descriptions in section 3.5
# ============================================================
idx = find_para('check_binding_status')
if idx:
    for j in range(idx, idx + 20):
        if j < len(doc.paragraphs) and 'agentLoginSessionId' in doc.paragraphs[j].text:
            if '提取' in doc.paragraphs[j].text or '自动提取' in doc.paragraphs[j].text:
                doc.paragraphs[j].text = doc.paragraphs[j].text.replace(
                    'agentLoginSessionId 由 XiaoYiAuthFilter 从 HTTP Header 自动提取',
                    '系统自动从请求 arguments 中提取设备标识（deviceInfo.sid 或 session.sessionId）')
                break

idx = find_para('bind_employee')
if idx:
    for j in range(idx, idx + 20):
        if j < len(doc.paragraphs) and 'agentLoginSessionId' in doc.paragraphs[j].text:
            if '提取' in doc.paragraphs[j].text or '自动提取' in doc.paragraphs[j].text:
                doc.paragraphs[j].text = doc.paragraphs[j].text.replace(
                    'agentLoginSessionId 由 XiaoYiAuthFilter 从 HTTP Header 自动提取',
                    '系统自动从请求 arguments 中提取设备标识（deviceInfo.sid 或 session.sessionId）')
                break

# ============================================================
# 3. Update section 3.3 XiaoYiAuthFilter - simplified
# ============================================================
idx = find_para('3.3 XiaoYiAuthFilter')
if idx:
    for j in range(idx, idx + 10):
        if j < len(doc.paragraphs) and '存入' in doc.paragraphs[j].text and 'ThreadLocal' in doc.paragraphs[j].text:
            doc.paragraphs[j].text = '从 HTTP Header 提取 appId/apiKey 做认证，记录所有请求头日志。不再提取 agentLoginSessionId（已改由 arguments 传递）。'
            break

# ============================================================
# 4. Update section 3.4 - add new note
# ============================================================
idx = find_para('3.4 MCP 服务端点')
if idx:
    for j in range(idx, idx + 20):
        if j < len(doc.paragraphs) and 'auth MCP' in doc.paragraphs[j].text and 'POST' in doc.paragraphs[j].text:
            # Update the auth MCP line
            doc.paragraphs[j].text = 'auth MCP：POST /admin3/xiaoyi-mcp/auth/message（授权/解授权，暂未使用，保留结构）'
            break

# ============================================================
# 5. Update section 4 (data flow) - change to sid/sessionId
# ============================================================
# 4.1
idx = find_para('4.1 用户已授权（正常流程）')
if idx:
    for j in range(idx, idx + 15):
        if j < len(doc.paragraphs):
            if 'ThreadLocal' in doc.paragraphs[j].text and 'request.setAttribute' in doc.paragraphs[j].text:
                doc.paragraphs[j].text = '5. XiaoYiAuthFilter 从 Header 提取 appId/apiKey 完成认证'
            if 'contextExtractor' in doc.paragraphs[j].text:
                doc.paragraphs[j].text = '6. MCP 工具 handler 从 arguments 中提取 deviceInfo.sid 或 session.sessionId 识别用户'

# 4.2
idx = find_para('4.2 用户未绑定员工ID（首次使用）')
if idx:
    for j in range(idx, idx + 10):
        if j < len(doc.paragraphs):
            if 'XiaoYiAuthFilter' in doc.paragraphs[j].text and 'HTTP Header' in doc.paragraphs[j].text:
                doc.paragraphs[j].text = '4. 服务端自动从 arguments 提取 sid/sessionId，XiaoYiBindingService 核查员工ID是否存在，创建或更新会话绑定关系'
                break

# ============================================================
# 6. Update section 5 (关于 agentLoginSessionId 的传递机制)
# ============================================================
idx = find_para('五、关于 agentLoginSessionId 的传递机制（重要）')
if idx:
    # Find the next heading after section 5
    end_5 = idx
    for j in range(idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].style.name.startswith('Heading 1') or doc.paragraphs[j].style.name.startswith('Heading 2'):
            end_5 = j - 1
            break
    
    # Clear existing bullets in section 5 and replace
    # Actually let's just update the first line and add new bullets
    # Find the last bullet in section 5
    last_bullet_5 = end_5
    for j in range(idx, end_5 + 1):
        if doc.paragraphs[j].style.name == 'List Bullet':
            last_bullet_5 = j
    
    # Update the title line
    doc.paragraphs[idx].text = '五、设备标识传递机制（v2 变更：从 agentLoginSessionId 改为 sid/sessionId）'
    
    # Update the first normal paragraph
    for j in range(idx + 1, end_5 + 1):
        if doc.paragraphs[j].style.name == 'Normal' and doc.paragraphs[j].text.strip():
            doc.paragraphs[j].text = '由于小艺平台目前未开通华为授权流程，无法获取 agentLoginSessionId。v2 改为使用小艺 MCP 请求 arguments 中携带的设备标识。'
            break
    
    # Clear old bullets by marking them
    bullets_to_clear = []
    for j in range(idx + 1, end_5 + 1):
        if doc.paragraphs[j].style.name == 'List Bullet':
            bullets_to_clear.append(j)
    
    # Replace first bullet with new content
    if bullets_to_clear:
        doc.paragraphs[bullets_to_clear[0]].text = '小艺每次 MCP 请求在 JSON-RPC body 的 arguments 中携带 deviceInfo.sid（设备唯一标识）和 session.sessionId（会话ID）'
    
    if len(bullets_to_clear) > 1:
        doc.paragraphs[bullets_to_clear[1]].text = '服务端 XiaoYiAuthService.resolveDeviceKey() 从 arguments 中提取，优先使用 sid'
    
    if len(bullets_to_clear) > 2:
        doc.paragraphs[bullets_to_clear[2]].text = 'XiaoYiUserSession 表新增 sid（唯一索引）和 session_id 列，作为用户标识'
    
    if len(bullets_to_clear) > 3:
        doc.paragraphs[bullets_to_clear[3]].text = '绑定流程：用户输入 xEmployeeId → 服务端将 sid/sessionId 与 User 关联并保存'
    
    # Clear remaining old bullets
    for k in range(4, len(bullets_to_clear)):
        doc.paragraphs[bullets_to_clear[k]].text = ''
    
    # Add new bullets after last_bullet
    ref_p = doc.paragraphs[last_bullet_5]
    add_bullet_after(ref_p, '后续同一设备/会话的请求，sid 或 sessionId 任意一个匹配即可识别用户')
    add_bullet_after(ref_p, 'agentLoginSessionId 字段保留，待华为授权开通后可启用')

# ============================================================
# 7. Update section 6 (工作流设计建议)
# ============================================================
idx = find_para('六、工作流设计建议')
if idx:
    for j in range(idx, idx + 5):
        if j < len(doc.paragraphs) and 'sessionId' in doc.paragraphs[j].text and '提取' in doc.paragraphs[j].text:
            doc.paragraphs[j].text = 'check_binding_status 和 bind_employee 的设备标识（sid/sessionId）由服务端自动从 arguments 中提取，工作流中所有 MCP 工具调用都不需要关心设备标识参数。建议工作流结构如下：'
            break

# ============================================================
# 8. Update table 1 (file list) - add new/changed files
# ============================================================
if len(doc.tables) >= 2:
    table = doc.tables[1]
    new_rows = [
        ['XiaoYiUserSession.java', 'infra/mcp/xiaoyi/XiaoYiUserSession.java', '已修改：新增 sid、session_id 字段，agentLoginSessionId 改为非必填'],
        ['XiaoYiUserSessionRepository.java', 'infra/mcp/xiaoyi/XiaoYiUserSessionRepository.java', '已修改：新增 findBySid、findBySessionId'],
        ['XiaoYiAuthService.java', 'infra/mcp/xiaoyi/XiaoYiAuthService.java', '已修改：新增 resolveDeviceKey/findSessionByDeviceKey，checkBindingStatus 改 deviceKey'],
        ['XiaoYiBindingService.java', 'infra/mcp/xiaoyi/XiaoYiBindingService.java', '已修改：bindEmployee 改接收 sid+sessionId'],
        ['XiaoYiBindingMcpConfig.java', 'infra/mcp/xiaoyi/XiaoYiBindingMcpConfig.java', '已修改：去掉 contextExtractor，handler 从 arguments 提取 sid/sessionId'],
        ['XiaoYiAuthFilter.java', 'infra/mcp/xiaoyi/XiaoYiAuthFilter.java', '已修改：简化，只保留认证和 header 日志'],
        ['init_userbindtest_data.py', 'init_userbindtest_data.py', '已修改：插入 sid 测试数据，含自动迁移逻辑'],
        ['test_check_binding.py', 'test_check_binding.py', '已修改：arguments 传 sid/sessionId，不再传 agentLoginSessionId header'],
        ['src/main/resources/data.sql', 'src/main/resources/data.sql', '已修改：追加 sid/session_id 迁移脚本和测试数据'],
    ]
    for row_data in new_rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = cell_text

# ============================================================
# 9. Update todo list
# ============================================================
idx = find_para('八、待办事项')
if idx:
    # Find existing todo items
    todo_items = []
    for j in range(idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].style.name == 'List Bullet':
            if doc.paragraphs[j].text.strip():
                todo_items.append(j)
        elif doc.paragraphs[j].style.name.startswith('Heading') and j > idx:
            break
    
    if todo_items:
        # Update existing todos
        doc.paragraphs[todo_items[0]].text = '[已完成] 开发 bind_employee MCP 工具（XiaoYiBindingMcpConfig + XiaoYiBindingService）'
    
    if len(todo_items) > 1:
        doc.paragraphs[todo_items[1]].text = '[已完成] 在 auth MCP 中注册 authorize/deauthorize 工具（替代独立授权回调接口）'
    if len(todo_items) > 2:
        doc.paragraphs[todo_items[2]].text = '[已完成] 将设备标识从 agentLoginSessionId 改为 sid/sessionId（v2 变更）'
    
    # Add new todos after the last one
    ref_p = doc.paragraphs[todo_items[-1]] if todo_items else doc.paragraphs[idx]
    add_bullet_after(ref_p, '测试 sid 绑定流程：check_binding_status → bind_employee（sid 方式）')
    add_bullet_after(ref_p, '在小艺开发者平台配置账号绑定设置（API URL 指向 auth MCP endpoint），留待后续')

# ============================================================
# Save
# ============================================================
doc.save(DOC_PATH)
print('文档更新完成！')
